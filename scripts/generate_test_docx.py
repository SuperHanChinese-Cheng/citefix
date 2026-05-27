"""Generate test .docx files with known AGLC4 errors for testing CiteFix.

Run: python scripts/generate_test_docx.py
Output: tests/fixtures/sample_essay.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_footnote_via_xml(document: Document, paragraph, footnote_text: str, index: int) -> None:
    """Add a footnote to a paragraph using direct XML manipulation.

    python-docx doesn't natively support footnotes, so we manipulate the XML directly.
    This is a simplified version — the real extractor.py will handle the full XML spec.
    """
    from lxml import etree

    # Word XML namespaces
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    # Add footnote reference in the paragraph
    run = paragraph.add_run()
    run_elem = run._element

    rpr = etree.SubElement(run_elem, f"{{{w_ns}}}rPr")
    etree.SubElement(rpr, f"{{{w_ns}}}rStyle", attrib={f"{{{w_ns}}}val": "FootnoteReference"})

    footnote_ref = etree.SubElement(run_elem, f"{{{w_ns}}}footnoteReference")
    footnote_ref.set(f"{{{w_ns}}}id", str(index))

    # TODO: Also add the footnote content to footnotes.xml
    # This is complex — for now, use a pre-built .docx with footnotes
    # The real test fixtures should be hand-crafted .docx files


def generate_sample_essay() -> None:
    """Generate a sample essay with intentional AGLC4 errors in footnotes.

    NOTE: Due to python-docx limitations with footnotes, it is MUCH easier to:
    1. Create the essay in Microsoft Word or LibreOffice
    2. Add footnotes manually with intentional errors
    3. Save as .docx to tests/fixtures/

    This script generates the BODY TEXT only. Add footnotes manually.

    Intended footnotes (with errors marked):

    FN 1: Mabo vs Queensland (No 2) [1992] 175 CLR 1 at p. 42.
           ERRORS: "vs" → "v", square brackets → round, "at p." → ","

    FN 2: Mabo vs Queensland (No 2) [1992] 175 CLR 1 at p. 55.
           ERRORS: same as FN 1, PLUS should be "Ibid 55." (consecutive same source)

    FN 3: Corporations Act 2001 (Cth), Section 180(1).
           ERRORS: comma before section, "Section" → "s", not italicised

    FN 4: McCutcheon, Jani, "The Vanishing Author" (2013) 36 UNSWLJ 915, at p 920.
           ERRORS: surname-first, double quotes, abbreviated journal, "at p"

    FN 5: Palmer v Ayres [2017] HCA 5 at para 31.
           ERRORS: "at para" → ", [31]"

    FN 6: Mabo vs Queensland (No 2) [1992] 175 CLR 1, 42
           ERRORS: should be "Mabo (n 1) 42.", repeated full cite, "vs", brackets, no full stop

    FN 7: Mark Leeming, "Authority to Decide" (Federation Press, 2nd edition, 2020) p 45.
           ERRORS: double quotes → no quotes (book title italicised), "edition" → "ed", "p 45" → "45"

    FN 8: Limitation Act 2005 (WA) section 14(1)
           ERRORS: not italicised, "section" → "s", no full stop

    FN 9: ibid
           ERRORS: should be "Ibid." (capitalised, italicised, full stop)

    FN 10: Palmer v Ayres [2017] HCA 5, [31].
            CORRECT — this one should pass with zero issues
    """
    output_dir = Path(__file__).parent.parent / "tests" / "fixtures"
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading("Sample Essay with AGLC4 Errors", level=1)
    doc.add_paragraph("This document is a test fixture for CiteFix.")
    doc.add_paragraph(
        "The High Court's landmark decision established the doctrine of native title "
        "in Australian law.[1] The principles were further elaborated upon.[2]"
    )
    doc.add_paragraph(
        "Directors must exercise their powers with care and diligence under the "
        "relevant statutory provisions.[3] Scholars have critiqued the framework.[4]"
    )
    doc.add_paragraph(
        "The Federal Court has applied these principles consistently.[5] "
        "As the High Court originally noted.[6]"
    )
    doc.add_paragraph(
        "Leading commentators have provided detailed analysis.[7] "
        "Western Australian limitation periods apply.[8] "
        "The same provision is relevant here.[9] "
        "The Federal Court decision is clear.[10]"
    )

    doc.add_heading("Instructions", level=2)
    doc.add_paragraph(
        "TO COMPLETE THIS TEST FIXTURE:\n"
        "1. Open this file in Word or LibreOffice\n"
        "2. Replace [1], [2], etc. with actual footnotes (References → Insert Footnote)\n"
        "3. Type the footnote text listed in generate_test_docx.py docstring\n"
        "4. Save and close\n"
        "5. The file is now ready for testing"
    )

    output_path = output_dir / "sample_essay_template.docx"
    doc.save(str(output_path))
    print(f"Generated: {output_path}")
    print("NOTE: You must manually add footnotes in Word/LibreOffice. See docstring for content.")


if __name__ == "__main__":
    generate_sample_essay()
